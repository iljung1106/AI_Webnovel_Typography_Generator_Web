from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "reports" / "웹소설_타이포그래피_중간보고서_초안.docx"
SCREENSHOT_DIR = ROOT / "docs" / "reports" / "tempimg"


ACCENT = "25536B"
MUTED = "6B7280"
LIGHT = "EAF1F4"
LIGHTER = "F7FAFB"
BORDER = "C7D2D9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def style_table(table, header_rows: int = 1, widths: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    if widths:
        set_col_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
        cant_split.set(qn("w:val"), "true")
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
                for run in p.runs:
                    run.font.name = "Malgun Gothic"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                    run.font.size = Pt(8.8)
            if r_idx < header_rows:
                set_cell_shading(cell, LIGHT)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(ACCENT)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHTER)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    style_table(table, widths=widths)
    doc.add_paragraph()
    return table


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "EEF7F5")
    set_cell_border(cell, "B7D5CF")
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Malgun Gothic"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r2.font.size = Pt(9)
    doc.add_paragraph()


def add_placeholder(doc: Document, label: str, detail: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "F3F4F6")
    set_cell_border(cell, "CBD5E1")
    set_cell_margins(cell, top=260, start=220, bottom=260, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(detail)
    r2.font.name = "Malgun Gothic"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()


def add_screenshot(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(14.7))

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_implementation_screenshots(doc: Document) -> None:
    screenshots = sorted(SCREENSHOT_DIR.glob("*.png"))
    captions = [
        "그림 2. 장르 선택 화면: 장르별 예시 타이포그래피를 보고 제작 방향을 선택",
        "그림 3. 표지 업로드 화면: 표지는 미리보기와 배치 참고용으로 선택 업로드",
        "그림 4. 제목 입력 화면: 한글 작품 제목을 입력하고 다음 단계로 진행",
        "그림 5. AI 제목 배치 화면: 초기 배치를 생성하고 위치·크기 조절 가능",
        "그림 6. 스타일 정리 화면: 사용자의 짧은 입력을 요소/스타일 키워드로 정리",
        "그림 7. 타이포 시안 생성/선택 화면: 3개 후보 슬롯과 생성 상태 표시",
        "그림 8. 효과 편집 화면: 선택한 타이포에 재질·빛·그림자 효과를 적용",
        "그림 9. 내보내기 화면: 최종 PNG 및 레이어 ZIP export 흐름",
    ]
    if not screenshots:
        add_placeholder(
            doc,
            "그림 2. 실제 Web/App 화면 캡처 삽입 위치",
            "메인, 장르 선택, 배치 편집, 생성 대기/시안 선택, 효과 편집, 내보내기",
        )
        return

    doc.add_heading("7.1 실제 실행 화면", level=2)
    doc.add_paragraph(
        "다음 화면들은 현재 로컬 Web/App에서 확인 가능한 실제 실행 화면이다. "
        "장르 선택부터 표지 업로드, 제목 입력, AI 배치, 스타일 정리, 시안 생성, 효과 편집, 내보내기까지 "
        "핵심 사용자 흐름이 단계적으로 연결되어 있음을 보여준다."
    )
    for index, image in enumerate(screenshots):
        caption = captions[index] if index < len(captions) else f"그림 {index + 2}. 실제 실행 화면"
        add_screenshot(doc, image, caption)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 22, ACCENT),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 11.5, ACCENT),
        ("Heading 3", 10.5, "374151"),
    ]:
        st = styles[style_name]
        st.font.name = "Malgun Gothic"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 6)
        st.paragraph_format.space_after = Pt(5)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("웹소설 타이포그래피 제작 웹앱 중간보고서 초안")
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc() -> None:
    doc = Document()
    style_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("서비스데이터사이언스\n중간보고서")
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(10)
    title.add_run("웹소설 제목용 AI 타이포그래피 제작 웹앱")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(24)
    run = sub.add_run(
        "작품 제목과 장르, 사용자가 입력한 스타일 방향을 바탕으로\n"
        "한글 제목 타이포그래피 시안을 생성하고 표지 위에서 편집·내보내기까지 지원하는 Web/App"
    )
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("374151")

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["팀명", "추후 입력"],
            ["팀원 및 역할", "추후 입력"],
            ["프로젝트명", "웹소설 제목용 AI 타이포그래피 제작 웹앱"],
            ["대상 사용자", "웹소설 작가 개인, 1인 창작자, 소규모 매니지먼트/출판사"],
            ["제출일", "2026년 5월 26일"],
            ["현재 개발 단계", "기획·아키텍처·핵심 워크플로우 구현 진행 / 화면 캡처 및 최신 구현 상태 보완 필요"],
        ],
        widths=[4.0, 12.0],
    )

    doc.add_page_break()

    doc.add_heading("목차", level=1)
    add_numbered(
        doc,
        [
            "프로젝트 개요",
            "고객 니즈 및 제품 목표",
            "Product Goal → AI/ML Framing",
            "서비스 구조: IA / User Flow / Wireflow",
            "데이터 및 저장 정책",
            "모델/AI 접근법 및 구현 전략",
            "현재 구현된 Web/App 기능",
            "수업 내용 반영 매핑",
            "성과 측정 계획",
            "현재 애로사항 및 Trouble Shooting 요청",
            "최종 발표까지의 개발 계획",
        ],
    )

    doc.add_heading("1. 프로젝트 개요", level=1)
    doc.add_paragraph(
        "본 프로젝트는 웹소설 표지 일러스트 자체를 생성하는 서비스가 아니라, 표지 위에 들어갈 "
        "제목 타이포그래피를 제작하는 AI 보조 Web/App이다. 사용자는 장르를 선택하고 작품 제목과 "
        "간단한 스타일 단어를 입력한 뒤, AI가 제안한 제목 배치와 흑백 타이포그래피 시안을 확인한다. "
        "이후 선택한 타이포그래피를 투명 배경 이미지로 변환하고, 브라우저 기반 효과 편집기에서 "
        "표지 위에 배치한 뒤 PNG 또는 레이어 파일로 내보낸다."
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["대상 사용자", "웹소설 작가 개인, 1인 창작자, 소규모 매니지먼트/출판사"],
            ["해결하려는 문제", "제목 타이포그래피 외주 비용과 제작 시간이 부담스럽고, 직접 제작에는 디자인 지식과 도구 숙련이 필요함"],
            ["핵심 가치", "낮은 비용으로 여러 타이포그래피 시안을 빠르게 생성하고, 표지 위에서 바로 확인·수정·내보내기 가능"],
            ["서비스 범위", "한글 제목 타이포그래피 생성, 배치 편집, 효과 적용, PNG/ZIP 내보내기"],
            ["제외 범위", "표지 일러스트 생성, 장기 보관형 디자인 툴, 팀 협업, PSD/SVG export, 인간 디자이너 리터칭"],
        ],
        widths=[3.2, 12.8],
    )

    doc.add_heading("2. 고객 니즈 및 제품 목표", level=1)
    doc.add_heading("2.1 목표 시장과 사용자", level=2)
    doc.add_paragraph(
        "시장조사 자료에서는 핵심 타겟을 아마추어 및 자유연재 작가, 중소형 출판사 또는 웹소설 "
        "에이전시, 표지 작업을 수행하는 디자이너/일러스트레이터로 구분하였다. 본 MVP에서는 "
        "초기 진입 장벽과 구매 의사결정 부담이 낮은 개인 작가와 소규모 매니지먼트를 우선 대상으로 설정한다."
    )
    add_table(
        doc,
        ["사용자군", "Pain Point", "서비스가 제공하는 가치"],
        [
            ["개인 웹소설 작가", "표지는 준비했지만 제목 로고 제작 비용과 시간이 부담됨", "저렴한 크레딧으로 즉시 여러 타이포 시안을 확인"],
            ["소규모 매니지먼트/출판사", "다수 작품의 표지 시안과 유료화 전 테스트용 타이포가 반복적으로 필요함", "제작 공정 단축과 대량 초안 확보"],
            ["디자이너/일러스트레이터", "타이포 초안과 스타일 레퍼런스를 빠르게 잡아야 함", "작업 전 아이디어 후보와 효과 방향을 빠르게 확보"],
        ],
        widths=[3.2, 6.0, 6.8],
    )

    doc.add_heading("2.2 Product Goal", level=2)
    add_note(
        doc,
        "Product Goal",
        "웹소설 작가와 소규모 제작자가 전문 타이포그래피 외주 전 단계에서, 작품 제목과 장르를 바탕으로 "
        "사용 가능한 한글 제목 타이포 시안을 빠르게 생성하고 표지 위에서 확인·수정·내보낼 수 있도록 한다.",
    )
    doc.add_paragraph(
        "제품 목표는 ‘AI를 활용한다’가 아니라 사용자가 실제 작업에 쓸 수 있는 제목 타이포그래피 결과물을 얻도록 돕는 것이다. "
        "AI는 제목 배치, 스타일 키워드 정리, 흑백 타이포그래피 생성에 사용되며, 사용자는 각 단계에서 결과를 확인하고 직접 수정할 수 있다."
    )

    doc.add_heading("3. Product Goal → AI/ML Framing", level=1)
    add_table(
        doc,
        ["구분", "정의"],
        [
            ["Input", "장르, 작품 제목, 사용자가 입력한 스타일 단어/문구, 선택적 표지 이미지"],
            ["Intermediate Output", "초기 제목 배치, 요소/스타일 키워드 목록, 흑백 타이포 후보, 투명 배경 타이포"],
            ["Final Output", "표지 위에 배치된 최종 PNG, 투명 배경 타이포 PNG, 고급 export용 레이어 PNG ZIP"],
            ["f(x) → y", "f(장르, 제목, 스타일 단어, 선택적 표지 이미지) → 타이포 후보 3개와 편집 가능한 표지 합성 결과"],
        ],
        widths=[3.2, 12.8],
    )
    doc.add_heading("3.1 AI/ML 기능의 위치", level=2)
    add_table(
        doc,
        ["단계", "AI/ML 또는 자동화 역할", "현재 구현 방향"],
        [
            ["제목 배치", "제목 길이와 장르를 바탕으로 초기 위치·크기·회전값 생성", "기존 LayoutModule 로직을 API/Worker에서 호출"],
            ["스타일 정리", "사용자 입력을 짧은 영어 요소/스타일 토큰으로 정리", "OpenRouter 기반 PromptGenerationModule 사용"],
            ["타이포 생성", "색·재질·명암 없는 순수 흑백 실루엣 벡터풍 후보 생성", "Comfy Cloud를 통해 3개 슬롯 배치 생성"],
            ["투명화", "흑백 결과를 그레이스케일/threshold 곡선으로 투명 배경 마스크화", "Python 이미지 처리 후 candidate와 함께 저장"],
            ["효과 적용", "선택한 투명 타이포 위에 glow, shadow, ray beam, material 효과 적용", "TypoEffector 로직을 브라우저 렌더링으로 재사용"],
        ],
        widths=[2.6, 7.0, 6.4],
    )
    add_note(
        doc,
        "표지 이미지 사용 범위",
        "업로드한 표지는 미리보기 배경, 색상/밝기 분석, 빛 방향 추정, 자동 배치 참고에만 사용한다. "
        "표지 이미지는 타이포그래피 생성 모델의 입력, 모델 학습, 홍보 사례, 공개 갤러리 용도로 사용하지 않는다.",
    )

    doc.add_heading("4. 서비스 구조: IA / User Flow / Wireflow", level=1)
    doc.add_paragraph(
        "사용자 흐름은 Canva식 자유 편집 툴보다 선형적인 제작 절차를 우선한다. 모든 단계에는 뒤로 가기와 다음으로 가기가 있으며, "
        "사용자는 자동 제안을 먼저 받고 필요한 경우에만 직접 조정한다."
    )
    add_table(
        doc,
        ["순서", "화면/단계", "사용자 행동", "시스템 처리"],
        [
            ["1", "랜딩/시작", "서비스 설명 확인 후 만들러 가기 선택", "로그인 상태 확인 및 새 프로젝트 draft 시작"],
            ["2", "장르 선택", "로맨스 판타지, 현대, 판타지, 무협, 힐링 중 선택", "장르별 프롬프트 프로필과 초기 스타일 방향 설정"],
            ["3", "표지 업로드/제목 입력", "표지를 선택적으로 업로드하고 제목 입력", "표지 미리보기, 색/빛/배치 분석 준비"],
            ["4", "AI 배치", "초기 제목 배치를 확인하고 드래그·크기·회전 수정", "LayoutModule 기반 초기 배치 생성"],
            ["5", "요소/스타일 입력", "짧은 단어 또는 문구 입력 후 AI 정리 결과 확인", "LLM이 간결한 영어 키워드 목록 생성"],
            ["6", "생성 대기/시안 선택", "크레딧 차감 확인 후 3개 후보 중 하나 선택", "Comfy Cloud 작업, 슬롯별 상태·실패·환불 처리"],
            ["7", "효과 편집", "프리셋과 세부 설정을 조절하고 표지 위 위치 조정", "투명화된 타이포에 TypoEffector 기반 효과 적용"],
            ["8", "내보내기", "기본 PNG 또는 고급 ZIP export 선택", "브라우저 렌더링 기반 파일 생성 및 저장"],
        ],
        widths=[1.2, 3.1, 5.7, 6.0],
    )
    add_placeholder(doc, "그림 1. 핵심 사용자 흐름 캡처 삽입 위치", "랜딩 → 장르 선택 → 제목 입력 → 생성 → 효과 편집 → 내보내기")

    doc.add_heading("5. 데이터 및 저장 정책", level=1)
    add_table(
        doc,
        ["데이터 항목", "현재/예정 상태", "활용 목적", "주의점"],
        [
            ["사용자 계정", "Supabase Google OAuth", "사용자별 프로젝트 분리", "서비스 role key는 서버/워커에만 보관"],
            ["프로젝트/버전", "Supabase Postgres", "작업 흐름 저장, 최근 작업 조회", "개인 계정 전용, 팀 작업 미지원"],
            ["표지 이미지", "Supabase Storage private bucket", "미리보기, 색/빛/배치 분석", "AI 생성 입력·학습·홍보 비활용 명시"],
            ["제목/스타일 입력", "ProjectVersion JSON", "배치·프롬프트·생성 입력", "민감 정보 로그 저장 최소화"],
            ["생성 후보", "candidate/transparent asset", "시안 선택 및 효과 편집", "짧은 보관 기간과 삭제 정책 필요"],
            ["크레딧 로그", "Release 3에서 ledger 구현", "생성/고급 export 차감 및 환불", "실결제 전까지는 placeholder 처리"],
        ],
        widths=[3.0, 3.6, 4.4, 5.0],
    )

    doc.add_heading("6. 모델/AI 접근법 및 구현 전략", level=1)
    add_table(
        doc,
        ["접근법", "설명", "장점", "한계/대응"],
        [
            ["Rule-based", "장르별 프리셋, 단순 배치 규칙, 효과 기본값", "빠르고 설명 가능함", "다양성이 낮아 AI 생성과 병행"],
            ["LLM Prompt Resolution", "사용자 표현을 생성 모델용 짧은 영어 토큰으로 정리", "입력 부담 감소", "묘사가 과도해지지 않도록 프롬프트 제한"],
            ["Diffusion/Comfy Cloud", "흑백 타이포 후보 3개 생성", "형태 다양성과 장식 요소 표현 가능", "한글 가독성·timeout·비용 관리 필요"],
            ["Image Processing", "흑백 시안을 투명 배경 마스크로 변환", "효과 편집에 필요한 안정적 중간 산출물 확보", "threshold와 anti-aliasing 품질 조정 필요"],
            ["Browser Rendering", "TypoEffector 기반 효과와 합성을 브라우저에서 처리", "서버 비용 절감, 즉시 편집 가능", "해상도 상한과 GPU 호환성 관리 필요"],
        ],
        widths=[3.1, 5.0, 3.5, 4.4],
    )
    doc.add_paragraph(
        "현재 전략은 완전 자동 디자인 생성이 아니라 hybrid approach이다. 기존에 작동하던 LayoutModule, PromptGenerationModule, "
        "ImageGenerationModule, TypoEffector 로직을 새 서비스 구조에 감싸거나 이식하여, 이전 프로토타입과 같거나 유사한 작동 방식과 성능을 유지한다."
    )

    doc.add_heading("7. 현재 구현된 Web/App 기능", level=1)
    doc.add_paragraph(
        "아래 표는 중간보고서 초안 기준의 구현 상태 정리 양식이다. 최종 제출 전에는 실제 실행 화면을 확인하여 상태를 "
        "‘실제 구현 / 부분 구현 / mock-up / hard-coded / manual / 미구현’ 중 하나로 확정하고 캡처를 삽입해야 한다."
    )
    add_table(
        doc,
        ["기능", "현재 구분", "연결 화면", "사용 데이터/모듈", "비고"],
        [
            ["랜딩/메인", "부분 구현", "메인 페이지", "Next.js", "서비스 진입, 로그인, 최근 작업 영역"],
            ["Google 로그인", "부분 구현", "상단/시작 흐름", "Supabase Auth", "배포 URL 설정과 OAuth redirect 최종 확인 필요"],
            ["프로젝트 저장/조회", "부분 구현", "최근 작업, 제작 화면", "Supabase DB", "사용자별 작업 분리 검증 필요"],
            ["장르/제목 입력", "실제 구현", "제작 workflow", "프론트 상태 + DB draft", "장르 목록은 로맨스 판타지/현대/판타지/무협/힐링"],
            ["AI 제목 배치", "부분 구현", "배치 탭", "LayoutModule/OpenRouter", "초기 중앙 정렬, 드래그/크기/회전 UX 보완 필요"],
            ["스타일 키워드 생성", "부분 구현", "스타일 입력 화면", "PromptGenerationModule/OpenRouter", "짧은 영어 토큰 중심으로 출력 제한"],
            ["타이포 후보 생성", "부분 구현", "생성/시안 선택 화면", "Comfy Cloud/ImageGenerationModule", "3개 슬롯, 실패/timeout/환불 규칙 필요"],
            ["투명화", "부분 구현", "생성 후 처리", "Pillow 기반 post-processing", "threshold/곡선 조정 진행"],
            ["효과 편집", "부분 구현", "효과 편집 화면", "TypoEffector port", "프리셋, 세부 설정, 빛 방향 GUI 보완 필요"],
            ["내보내기", "부분 구현", "내보내기 화면", "Browser canvas/ZIP", "기본 PNG와 레이어 ZIP 완성도 점검 필요"],
            ["결제/크레딧", "미구현/placeholder", "확인 모달", "Credit ledger 예정", "Release 3 범위"],
            ["실서버 배포", "미구현", "Vercel/Render/Supabase", "배포 환경변수", "최종 발표 전 production-like URL 확보 필요"],
        ],
        widths=[2.5, 2.3, 2.7, 4.0, 4.5],
    )
    add_implementation_screenshots(doc)

    doc.add_heading("8. 수업 내용 반영 매핑", level=1)
    add_table(
        doc,
        ["수업/실습 내용", "우리 프로젝트의 산출물", "반영 방식"],
        [
            ["고객 니즈 파악", "대상 사용자와 pain point", "개인 작가와 소규모 제작자의 비용·시간·디자인 지식 부족 문제 정의"],
            ["Product Goal", "서비스 목표 문장", "사용 가능한 제목 타이포 시안을 빠르게 생성·수정·내보내는 목표로 정리"],
            ["AI/ML Framing", "input-output 구조", "장르/제목/스타일/표지 분석 정보를 입력으로 후보 타이포와 합성 결과를 출력"],
            ["데이터 검토", "DB/Storage/보관 정책", "프로젝트, 버전, 작업, 슬롯, asset, credit ledger 구조 설계"],
            ["ML 접근법 선택", "Hybrid pipeline", "LLM, diffusion, 이미지 처리, 브라우저 렌더링을 각 역할에 맞게 분리"],
            ["Business Metric", "다운로드/export 전환율", "사용자가 실제 산출물을 얻는 행동을 핵심 성공 지표로 설정"],
            ["Model Metric", "가독성, 장르 적합성, 선택률, 실패율", "AI 결과 품질과 사용자 행동 지표를 연결"],
            ["IA/User Flow", "선형 workflow", "장르 선택부터 내보내기까지 단계별 확인 구조로 설계"],
            ["Vibe Coding", "Next.js/FastAPI/Worker 구현", "현재 Web/App skeleton과 핵심 workflow를 실제 코드로 구현 진행"],
            ["Iteration Plan", "Release 1~3", "생성 → 효과/export → 결제/운영 순서로 범위 분리"],
        ],
        widths=[3.2, 4.0, 8.8],
    )

    doc.add_heading("9. 성과 측정 계획", level=1)
    doc.add_heading("9.1 Business Metric", level=2)
    add_bullets(
        doc,
        [
            "생성 시작률: 제작 화면 진입 사용자 중 크레딧 확인 후 생성 요청까지 진행한 비율",
            "생성 완료율: 생성 batch 중 1개 이상 성공 후보를 받은 비율",
            "시안 선택률: 생성 완료 후 후보 중 하나를 선택한 비율",
            "Export 전환율: 편집 화면 진입 후 최종 PNG 또는 ZIP을 다운로드한 비율",
            "재방문/재작업률: 최근 작업에서 다시 편집하거나 새 프로젝트를 만든 사용자 비율",
        ],
    )
    doc.add_heading("9.2 Model / AI Quality Metric", level=2)
    add_bullets(
        doc,
        [
            "제목 가독성: 한글 제목을 사용자가 쉽게 읽을 수 있는지",
            "장르 적합성: 장르에 맞는 장식 요소와 실루엣이 반영되었는지",
            "시안 선택률: 3개 후보 중 최소 1개가 선택되는 비율",
            "수정량: 사용자가 후보 선택 후 위치·크기·효과를 얼마나 수정하는지",
            "생성 실패율과 latency: timeout, postprocess 실패, Comfy Cloud 실패 비율과 대기 시간",
        ],
    )
    doc.add_paragraph(
        "Business Metric과 Model Metric은 export 전환율을 중심으로 연결된다. 생성 품질이 높아지면 후보 선택률이 올라가고, "
        "후보 선택 이후 과도한 수정을 하지 않아도 최종 내보내기까지 이어질 가능성이 높아진다."
    )

    doc.add_heading("10. 현재 애로사항 및 Trouble Shooting 요청", level=1)
    add_table(
        doc,
        ["문제 유형", "현재 문제", "시도한 해결 방법", "필요한 도움"],
        [
            ["모델/AI", "생성형 모델이 색·텍스처·명암을 포함하거나 한글 가독성을 해칠 수 있음", "프롬프트에 순수 검은색 실루엣, white glow, 불필요한 묘사 제한을 명시", "가독성 유지 프롬프트와 평가 방식 피드백"],
            ["이미지 처리", "흑백 후보를 투명 배경으로 변환할 때 경계 품질과 투명 범위 조정 필요", "threshold와 곡선 조정, anti-aliasing 유지", "품질 기준과 자동 검증 방법 조언"],
            ["UX/UI", "편집 단계가 많아 보일 수 있고 드래그/크기/회전 UX가 중요함", "선형 workflow와 자동 기본값, 직접 수정 가능 구조로 설계", "실사용자 흐름 검토"],
            ["성능", "효과 설정 변경이나 위치 이동 시 렌더링 비용이 커질 수 있음", "템플릿/세부 설정 변경 시에만 debounced render, 위치 이동은 캐시 이미지 합성", "브라우저 렌더링 최적화 조언"],
            ["배포", "프론트/API/Worker/Supabase/OAuth 환경변수를 각각 맞춰야 함", "Vercel + Render + Supabase 구조로 분리", "배포 troubleshooting 및 보안 설정 검토"],
            ["법률/정책", "상업 사용, 저작권, 업로드 이미지 비활용, 책임 범위 정리가 필요함", "비활용 원칙과 사용자 입력 책임 원칙 문서화", "약관·개인정보 처리방침 법률 검토"],
        ],
        widths=[2.5, 5.0, 4.5, 4.0],
    )

    doc.add_heading("11. 최종 발표까지의 개발 계획", level=1)
    add_table(
        doc,
        ["우선순위", "개발 항목", "목표", "비고"],
        [
            ["1", "실제 사용자 흐름 검증", "랜딩 → 제작 → 생성 → 효과 → export를 한 번에 시연 가능하게 정리", "필수"],
            ["2", "로그인/사용자별 작업 분리", "Google 로그인과 최근 작업 조회가 계정별로 분리되도록 검증", "필수"],
            ["3", "생성 batch 안정화", "3개 슬롯 상태, 실패/timeout, 투명화 결과 저장", "필수"],
            ["4", "효과 편집 완성도 보완", "TypoEffector 효과, 프리셋, 세부 설정, 빛 방향 GUI 정리", "필수"],
            ["5", "내보내기 기능", "기본 PNG와 투명 타이포 PNG, 가능하면 레이어 ZIP 제공", "필수/확장"],
            ["6", "배포", "Vercel, Render, Supabase 환경에서 production-like URL 확보", "최종 발표 전"],
            ["7", "발표자료와 데모 시나리오", "10분 발표 흐름에 맞춰 화면 캡처와 실제 시연 순서 정리", "필수"],
            ["8", "크레딧/결제", "초기에는 placeholder, 실제 결제는 Release 3로 분리", "범위 조정"],
        ],
        widths=[1.7, 4.0, 7.2, 3.1],
    )
    add_note(
        doc,
        "최종 발표 범위",
        "반드시 보여줄 것은 아이디어 설명이 아니라 실제 사용자 흐름이다. 발표에서는 업로드/제목 입력/AI 배치/스타일 정리/3개 시안 생성/"
        "후보 선택/효과 편집/export 순서로 데모를 구성한다. 구현되지 않은 기능은 mock-up, placeholder, 미구현으로 명확히 구분한다.",
    )

    add_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
